from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from spellchecker import SpellChecker
import re
import os
import json
import uuid

import google.generativeai as genai
from docx import Document
from reportlab.pdfgen import canvas

app = Flask(__name__)
CORS(app)

# ----------------------------
# ✅ Spell Checker Setup
# ----------------------------
spell = SpellChecker()
custom_dictionary = set()

# ----------------------------
# ✅ Gemini Setup
# ----------------------------
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

# ----------------------------
# ✅ Export Folder Setup
# ----------------------------
EXPORT_DIR = os.path.join(os.path.dirname(__file__), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


def extract_words(text: str):
    """Extract words only (ignores punctuation)."""
    return re.findall(r"[A-Za-z']+", text)


def get_supported_gemini_model():
    """Return first Gemini model that supports generateContent."""
    try:
        for m in genai.list_models():
            if "generateContent" in m.supported_generation_methods:
                return m.name
        return None
    except Exception:
        return None


def clean_gemini_json(ai_text: str):
    """Remove markdown code blocks if Gemini returns them."""
    ai_text = ai_text.strip()
    ai_text = ai_text.replace("```json", "").replace("```", "").strip()
    return ai_text


@app.route("/", methods=["GET"])
def home():
    return jsonify({
        "message": "✅ Spell Checker API running (NO grammar checker)",
        "endpoints": {
            "POST /check": "Spell check",
            "POST /tone-detect": "Tone detection (Gemini)",
            "POST /ai-improve": "AI rewriting (Gemini)",
            "POST /export/pdf": "Export to PDF",
            "POST /export/docx": "Export to DOCX",
            "POST /add-word": "Add custom word",
            "POST /reset-dictionary": "Reset custom dictionary"
        }
    })


# ----------------------------
# ✅ Custom Dictionary
# ----------------------------
@app.route("/add-word", methods=["POST"])
def add_word():
    data = request.get_json() or {}
    word = data.get("word", "").strip().lower()

    if not word:
        return jsonify({"error": "Word is required"}), 400

    custom_dictionary.add(word)
    return jsonify({"message": f"✅ '{word}' added to custom dictionary"})


@app.route("/reset-dictionary", methods=["POST"])
def reset_dictionary():
    custom_dictionary.clear()
    return jsonify({"message": "✅ Custom dictionary cleared"})


# ----------------------------
# ✅ Spell Check Endpoint
# ----------------------------
@app.route("/check", methods=["POST"])
def check_spelling():
    data = request.get_json() or {}
    text = data.get("text", "")

    if not text.strip():
        return jsonify({"error": "Text is required"}), 400

    words = extract_words(text)
    words_lower = [w.lower() for w in words]

    to_check = [w for w in words_lower if w not in custom_dictionary]
    misspelled = list(spell.unknown(to_check))

    mistakes = []
    for w in misspelled:
        cand = spell.candidates(w)
        suggestions = list(cand) if cand else []
        mistakes.append({
            "word": w,
            "suggestions": suggestions[:6]
        })

    corrected_text = text
    for w in misspelled:
        best = spell.correction(w)
        if best:
            corrected_text = re.sub(
                rf"\b{re.escape(w)}\b",
                best,
                corrected_text,
                flags=re.IGNORECASE
            )

    total_words = len(words)
    wrong_count = len(misspelled)
    accuracy = 100 if total_words == 0 else round(((total_words - wrong_count) / total_words) * 100, 2)

    return jsonify({
        "total_words": total_words,
        "wrong_words_count": wrong_count,
        "accuracy": accuracy,
        "mistakes": mistakes,
        "corrected_text": corrected_text
    })


# ----------------------------
# ✅ Tone Detector (Gemini)
# ----------------------------
@app.route("/tone-detect", methods=["POST"])
def tone_detect():
    if not GEMINI_API_KEY:
        return jsonify({"error": "Gemini API Key not set in server environment"}), 500

    data = request.get_json() or {}
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "Text is required"}), 400

    model_name = get_supported_gemini_model()
    if not model_name:
        return jsonify({"error": "No Gemini model available for generateContent"}), 500

    try:
        model = genai.GenerativeModel(model_name)

        prompt = f"""
You are a Tone Detection AI.

Detect the tone of this text and return ONLY JSON (no markdown).
Possible tone values:
Professional, Friendly, Formal, Informal, Angry, Sad, Confident, Neutral.

Text:
{text}

Return JSON like:
{{
  "tone": "Professional",
  "confidence": "High",
  "suggestion": "To make it more professional, avoid slang and use formal words."
}}
"""

        response = model.generate_content(prompt)
        ai_text = response.text if response and response.text else ""
        ai_text = clean_gemini_json(ai_text)

        parsed = json.loads(ai_text)
        parsed["model_used"] = model_name
        return jsonify(parsed)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ----------------------------
# ✅ AI Improve Writing (Gemini)
# ----------------------------
@app.route("/ai-improve", methods=["POST"])
def ai_improve():
    if not GEMINI_API_KEY:
        return jsonify({"error": "Gemini API Key not set in server environment"}), 500

    data = request.get_json() or {}
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "Text is required"}), 400

    model_name = get_supported_gemini_model()
    if not model_name:
        return jsonify({"error": "No Gemini model available for generateContent"}), 500

    try:
        model = genai.GenerativeModel(model_name)

        prompt = f"""
You are an AI writing assistant.

Rewrite the text in a professional way.
Return ONLY JSON (no markdown).

Text:
{text}

Return JSON like:
{{
  "professional_version": "..."
}}
"""

        response = model.generate_content(prompt)
        ai_text = response.text if response and response.text else ""
        ai_text = clean_gemini_json(ai_text)

        parsed = json.loads(ai_text)
        parsed["model_used"] = model_name
        return jsonify(parsed)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ----------------------------
# ✅ Export DOCX
# ----------------------------
@app.route("/export/docx", methods=["POST"])
def export_docx():
    data = request.get_json() or {}
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "Text is required"}), 400

    try:
        file_id = str(uuid.uuid4())[:8]
        filename = f"spell_checker_export_{file_id}.docx"
        filepath = os.path.join(EXPORT_DIR, filename)

        doc = Document()
        doc.add_heading("Spell Checker Export", level=1)
        doc.add_paragraph(text)
        doc.save(filepath)

        return send_file(filepath, as_attachment=True, download_name=filename)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ----------------------------
# ✅ Export PDF
# ----------------------------
@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    data = request.get_json() or {}
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "Text is required"}), 400

    try:
        file_id = str(uuid.uuid4())[:8]
        filename = f"spell_checker_export_{file_id}.pdf"
        filepath = os.path.join(EXPORT_DIR, filename)

        c = canvas.Canvas(filepath)
        c.setTitle("Spell Checker Export")
        c.setFont("Helvetica", 12)

        y = 800
        c.drawString(50, y, "Spell Checker Export")
        y -= 30

        for line in text.split("\n"):
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 12)
                y = 800

            # keep line safe width
            c.drawString(50, y, line[:120])
            y -= 18

        c.save()

        return send_file(filepath, as_attachment=True, download_name=filename)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)
